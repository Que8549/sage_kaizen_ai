"""
document_parser.py

Extracts plain text from uploaded files (text, code, Office documents) so they
can be injected into the LLM message context as named document blocks.

Supported formats
-----------------
Plain text / code  : .txt .md .csv .json .yaml .yml .toml .xml .html .htm
                     .log .cfg .ini .env
Code files         : .py .cs .js .ts .go .rs .java .cpp .c .h .hpp .sql
                     .sh .ps1 .rb .php .swift .kt .r .lua .zig
Word documents     : .docx  (python-docx; includes paragraphs + tables)
Excel workbooks    : .xlsx  (openpyxl; sheet-by-sheet row extraction)

Character limits
----------------
PER_DOCUMENT_CHAR_LIMIT : hard cap per individual file (~12 K tokens at 4 chars/token).
TOTAL_CHAR_LIMIT         : cap across all documents in one turn (~37 K tokens).

Both caps protect ARCHITECT's 128K context from accidental over-fill.  A truncation
notice is appended so the model knows the content was cut.
"""
from __future__ import annotations

import io
from dataclasses import dataclass

from sk_logging import get_logger

_LOG = get_logger("sage_kaizen.document_parser")

# ---------------------------------------------------------------------------
# Character limits
# ---------------------------------------------------------------------------

# Per-document cap: 50 000 chars ≈ 12 500 tokens (conservative for 4-char avg).
# Keeps a single large file from consuming the entire ARCHITECT context window.
PER_DOCUMENT_CHAR_LIMIT: int = 50_000

# Total cap across ALL documents in one turn: 150 000 chars ≈ 37 500 tokens.
# ARCHITECT context is 128K tokens; this leaves ~90K tokens for system prompt,
# conversation history, RAG/wiki/search context, and the model's response.
TOTAL_CHAR_LIMIT: int = 150_000

# Truncation notice appended when content is cut.
_TRUNCATION_NOTICE = "\n\n[... content truncated — file exceeded per-document character limit ...]"

# ---------------------------------------------------------------------------
# Extension lookup tables
# ---------------------------------------------------------------------------

# Extensions whose raw bytes decode directly to UTF-8 text.
# Mapping: extension → human-readable type label (used in the <document> tag).
_PLAIN_TEXT_TYPES: dict[str, str] = {
    # Generic text
    "txt":  "text",
    "md":   "markdown",
    "csv":  "csv",
    "log":  "log",
    "cfg":  "config",
    "ini":  "config",
    "env":  "env",
    "toml": "toml",
    "yaml": "yaml",
    "yml":  "yaml",
    # Data / markup
    "json": "json",
    "xml":  "xml",
    "html": "html",
    "htm":  "html",
    # Python
    "py":   "python",
    # .NET / C
    "cs":   "csharp",
    "c":    "c",
    "cpp":  "cpp",
    "h":    "c-header",
    "hpp":  "cpp-header",
    # Web / scripting
    "js":   "javascript",
    "ts":   "typescript",
    "rb":   "ruby",
    "php":  "php",
    "lua":  "lua",
    # Systems / compiled
    "go":   "go",
    "rs":   "rust",
    "swift":"swift",
    "kt":   "kotlin",
    "zig":  "zig",
    # JVM
    "java": "java",
    "r":    "r",
    # Data / query
    "sql":  "sql",
    # Shell
    "sh":   "bash",
    "ps1":  "powershell",
}

# Office formats handled with dedicated parsers.
_OFFICE_TYPES: dict[str, str] = {
    "docx": "word-document",
    "xlsx": "excel-workbook",
}

# Full set of accepted extensions (used by the UI to build the accept list).
ACCEPTED_EXTENSIONS: frozenset[str] = frozenset(_PLAIN_TEXT_TYPES) | frozenset(_OFFICE_TYPES)


# ---------------------------------------------------------------------------
# DocumentAttachment dataclass
# ---------------------------------------------------------------------------

@dataclass(frozen=True)
class DocumentAttachment:
    """
    A single text-based document attached to a chat turn.

    Unlike MediaAttachment (binary audio/image/video), DocumentAttachment stores
    extracted plain text — no base64 encoding needed.

    Attributes
    ----------
    filename   : Original filename as uploaded (e.g. ``"main.py"``).
    content    : Extracted plain text (possibly truncated — see PER_DOCUMENT_CHAR_LIMIT).
    doc_type   : Human-readable format label (``"python"``, ``"word-document"``, etc.).
    char_count : Character count of the *extracted* content (after any truncation).
    truncated  : True when the file was longer than PER_DOCUMENT_CHAR_LIMIT.
    """
    filename:   str
    content:    str
    doc_type:   str
    char_count: int
    truncated:  bool = False


# ---------------------------------------------------------------------------
# Internal extractors
# ---------------------------------------------------------------------------

def _extract_plain_text(raw: bytes, filename: str) -> str:
    """
    Decode raw bytes as UTF-8 text, falling back to latin-1 on decode errors.

    latin-1 is a lossless fallback for any byte sequence (every byte 0x00–0xFF
    maps to a valid code point), so the content is never silently lost.
    """
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        _LOG.debug("document_parser: %r is not UTF-8; falling back to latin-1", filename)
        return raw.decode("latin-1")


def _extract_docx(raw: bytes, filename: str) -> str:
    """
    Extract text from a .docx file using python-docx.

    Extracts paragraph text in document order, then appends table contents
    formatted as pipe-delimited rows.  Empty paragraphs are collapsed.
    """
    try:
        import docx  # python-docx

        doc = docx.Document(io.BytesIO(raw))
        lines: list[str] = []

        # Body paragraphs (preserves heading/body order)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)

        # Tables — formatted as pipe-delimited rows for LLM readability
        for table in doc.tables:
            lines.append("")  # blank separator before table
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                # Deduplicate merged cells (python-docx repeats cell text for merged spans)
                deduped: list[str] = []
                prev = object()
                for cell in cells:
                    if cell != prev:
                        deduped.append(cell)
                        prev = cell
                row_text = " | ".join(deduped)
                if row_text.strip():
                    lines.append(row_text)

        return "\n".join(lines)

    except Exception as exc:
        _LOG.warning("document_parser: failed to extract .docx from %r: %s", filename, exc)
        return f"[Error extracting document content: {exc}]"


def _extract_xlsx(raw: bytes, filename: str) -> str:
    """
    Extract text from a .xlsx file using openpyxl.

    Each sheet is rendered as a labelled block.  Rows are pipe-delimited.
    Fully empty rows are skipped.  read_only + data_only avoids formula
    evaluation and keeps memory usage low for large workbooks.
    """
    try:
        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
        sections: list[str] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_text: list[str] = []

            for row in ws.iter_rows(values_only=True):
                cells = [str(cell) if cell is not None else "" for cell in row]
                # Skip rows that are entirely empty or whitespace
                if any(c.strip() for c in cells):
                    rows_text.append(" | ".join(cells))

            if rows_text:
                sections.append(f"### Sheet: {sheet_name}")
                sections.extend(rows_text)

        wb.close()
        return "\n".join(sections)

    except Exception as exc:
        _LOG.warning("document_parser: failed to extract .xlsx from %r: %s", filename, exc)
        return f"[Error extracting workbook content: {exc}]"


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def _file_extension(filename: str) -> str:
    """Return the lowercase extension without the leading dot, or ''."""
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def is_supported(filename: str) -> bool:
    """Return True when this file type can be parsed by document_parser."""
    return _file_extension(filename) in ACCEPTED_EXTENSIONS


def parse_document(raw: bytes, filename: str) -> DocumentAttachment | None:
    """
    Extract plain text from *raw* bytes and return a DocumentAttachment.

    Returns None when the file type is unsupported.  Truncates content that
    exceeds PER_DOCUMENT_CHAR_LIMIT and sets ``truncated=True``.

    Parameters
    ----------
    raw      : Raw bytes of the uploaded file.
    filename : Original filename (used for extension detection and labelling).
    """
    ext = _file_extension(filename)

    if ext in _PLAIN_TEXT_TYPES:
        doc_type = _PLAIN_TEXT_TYPES[ext]
        content  = _extract_plain_text(raw, filename)
    elif ext == "docx":
        doc_type = _OFFICE_TYPES["docx"]
        content  = _extract_docx(raw, filename)
    elif ext == "xlsx":
        doc_type = _OFFICE_TYPES["xlsx"]
        content  = _extract_xlsx(raw, filename)
    else:
        _LOG.debug("document_parser: unsupported extension %r for %r", ext, filename)
        return None

    truncated = len(content) > PER_DOCUMENT_CHAR_LIMIT
    if truncated:
        content = content[:PER_DOCUMENT_CHAR_LIMIT] + _TRUNCATION_NOTICE
        _LOG.info(
            "document_parser: %r truncated to %d chars (original length exceeded limit)",
            filename, PER_DOCUMENT_CHAR_LIMIT,
        )

    return DocumentAttachment(
        filename=filename,
        content=content,
        doc_type=doc_type,
        char_count=len(content),
        truncated=truncated,
    )


def format_document_context(attachments: tuple["DocumentAttachment", ...]) -> str:
    """
    Render all DocumentAttachments as a formatted context block for injection
    into the LLM user message.

    Each document is wrapped in an XML-style ``<document>`` tag carrying
    filename, type, and character count.  This format is well-recognised by
    modern instruction-tuned LLMs (consistent with Anthropic's own prompting
    guidelines) and clearly delimits document boundaries in long contexts.

    The total character count across all documents is capped at TOTAL_CHAR_LIMIT.
    If the aggregate exceeds the cap, later documents are truncated with a notice.

    Returns an empty string when *attachments* is empty.
    """
    if not attachments:
        return ""

    parts: list[str] = []
    total_chars = 0

    for att in attachments:
        remaining = TOTAL_CHAR_LIMIT - total_chars
        if remaining <= 0:
            parts.append(
                f"\n[Document '{att.filename}' omitted — total document budget exhausted.]"
            )
            continue

        content = att.content
        if len(content) > remaining:
            content = content[:remaining] + _TRUNCATION_NOTICE
            _LOG.info(
                "format_document_context: '%s' further truncated to fit total budget (%d chars remaining)",
                att.filename, remaining,
            )

        total_chars += len(content)
        parts.append(
            f'<document filename="{att.filename}" type="{att.doc_type}" chars="{att.char_count}">\n'
            f"{content}\n"
            f"</document>"
        )

    return "\n\n".join(parts)
